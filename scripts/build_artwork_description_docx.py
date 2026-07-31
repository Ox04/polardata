from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "작품설명서_북극의기억층.docx"
IMAGE = ROOT / "outputs/final/cover-v7.png"
Q12_SOURCE = ROOT / "q12-description.md"

FONT = "Noto Sans CJK KR"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "152B3C"
MUTED = "5F7380"
LIGHT_FILL = "F4F6F9"
ICE_BLUE = "DCEFF4"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BULLET_NUM_ID: int | None = None


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell,
    *,
    top: int = 100,
    start: int = 120,
    bottom: int = 100,
    end: int = 120,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("Table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet")
    level.append(num_format)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    level.append(suffix)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)

    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "280")
    paragraph_properties.append(indentation)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    numbering.append(num)
    return num_id


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))
    set_run_font(run, size=9, color=MUTED)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)

    if "Source" not in [style.name for style in document.styles]:
        source = document.styles.add_style("Source", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = document.styles["Source"]
    source.font.name = FONT
    source._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    source.font.size = Pt(9)
    source.font.color.rgb = RGBColor.from_string(MUTED)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)


def configure_section(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("2026 극지 빅데이터-인공지능 활용 경진대회  |  작품 설명서")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("북극의 기억층  |  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(paragraph)


def add_title_block(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("북극의 기억층")
    set_run_font(run, size=27, bold=True, color=INK)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("Layers of Arctic Memory")
    set_run_font(run, size=14, color=BLUE)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(
        "1984–2024 북극 해빙 연령 구조를 기록한 3D 데이터 시각화"
    )
    set_run_font(run, size=10.5, bold=True, color=MUTED)

    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1200, 3480, 1200, 3480])
    metadata = (
        ("출품 부문", "데이터 아트", "작품 형식", "영상 · MP4"),
        ("팀명", "________________", "대표자명", "________________"),
    )
    for row_index, values in enumerate(metadata):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            is_label = column_index % 2 == 0
            set_run_font(
                run,
                size=9.5 if is_label else 10,
                bold=is_label,
                color=DARK_BLUE if is_label else INK,
            )
            if is_label:
                set_cell_shading(cell, LIGHT_FILL)

    document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(IMAGE), width=Inches(6.5))
    drawing = run._r.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", "1984년과 2024년 북극 해빙 연령을 같은 좌표에서 비교한 장면")
        drawing[0].set("title", "북극의 기억층 대표 장면")

    caption = document.add_paragraph(style="Caption")
    caption.add_run(
        "그림 1. 같은 제11주, 같은 좌표에서 비교한 1984년과 2024년의 해빙 연령 구조"
    )

    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(callout, ICE_BLUE)
    run = callout.add_run(
        "핵심 메시지  |  겨울은 얼음의 넓이를 되돌릴 수 있지만, "
        "사라진 얼음의 나이까지 한 번에 되돌리지는 못한다."
    )
    set_run_font(run, size=11, bold=True, color=DARK_BLUE)


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)
    return paragraph


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_bullet(document: Document, text: str) -> None:
    if BULLET_NUM_ID is None:
        raise RuntimeError("Bullet numbering was not initialized")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(BULLET_NUM_ID))
    num_properties.extend((level, num_id))
    paragraph_properties.append(num_properties)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_q12_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_data_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [2760, 1980, 1980, 2640])
    headers = ("지표", "1984년", "2024년", "변화")
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = ""
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])

    rows = (
        ("전체 해빙 범위", "약 1,067만 km²", "약 1,083만 km²", "약 1.5% 증가"),
        ("2년차 이상 해빙 비율", "44.4%", "22.3%", "22.1%p 감소"),
        ("5년차 이상 해빙 격자", "15,054개", "1,103개", "약 92.7% 감소"),
    )
    for row_values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_values):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=9.5, color=INK, bold=index == 0)
    set_table_geometry(table, [2760, 1980, 1980, 2640])


def read_q12_paragraphs() -> list[str]:
    lines = Q12_SOURCE.read_text(encoding="utf-8").splitlines()
    paragraphs: list[str] = []
    buffer: list[str] = []
    for line in lines:
        if line.startswith("#"):
            continue
        if not line.strip():
            if buffer:
                paragraphs.append(" ".join(buffer))
                buffer = []
            continue
        buffer.append(line.strip())
    if buffer:
        paragraphs.append(" ".join(buffer))
    return paragraphs


def build_document() -> None:
    global BULLET_NUM_ID
    document = Document()
    configure_styles(document)
    configure_section(document)
    BULLET_NUM_ID = create_bullet_numbering(document)
    add_title_block(document)

    add_heading(document, "1. 작품 개요")
    add_body(
        document,
        "《북극의 기억층》은 1984년부터 2024년까지 북극 해빙의 연령 구조를 "
        "색·높이·시간으로 변환한 44.5초 데이터 시각화 영상이다. 해빙의 바깥 "
        "경계보다 같은 범위 안에서 여러 해 살아남은 얼음이 어떻게 사라지는지 보여준다.",
    )

    add_heading(document, "2. 활용 데이터")
    add_body(
        document,
        "NASA Earthdata에서 제공하는 NSIDC의 EASE-Grid Sea Ice Age, "
        "Version 4.1(NSIDC-0611)을 활용했다. 해빙 이동 벡터로 얼음 입자를 주 "
        "단위로 추적해 격자별 최고 연령을 기록한 자료다.",
    )
    add_bullet(document, "분석 기간: 1984–2024년, 총 41개 연도")
    add_bullet(document, "비교 시점: 각 연도의 제11주")
    add_bullet(
        document,
        "연령 범주: 1년차, 2년차, 3년차, 4년차, 5년차 이상",
    )
    add_bullet(
        document,
        "핵심 지표: 2년차 이상 격자 수 ÷ 전체 해빙 격자 수",
    )
    source = document.add_paragraph(style="Source")
    source.add_run(
        "데이터 출처: Tschudi et al., EASE-Grid Sea Ice Age, Version 4.1, "
        "NASA NSIDC DAAC, https://doi.org/10.5067/UTAV7490FEPB"
    )

    add_heading(document, "3. 작품 기획 의도")
    add_body(
        document,
        "북극 해빙은 겨울마다 다시 넓어지지만 과거와 같은 상태로 회복되었다고 "
        "단정할 수는 없다. 새로 언 얼음과 여러 번의 여름을 견딘 얼음은 물리적 "
        "성질과 지속성이 다르기 때문이다. 따라서 ‘얼음이 얼마나 넓은가’보다 "
        "‘얼마나 오래 살아남았는가’를 보여주는 것을 출발점으로 삼았다.",
    )

    add_heading(document, "4. 데이터 처리 및 표현 방법")
    add_bullet(
        document,
        "원본 격자에서 육지·바다·결측값을 분리하고 해빙 연령을 다섯 범주로 재분류했다.",
    )
    add_bullet(
        document,
        "색은 짙은 청색에서 밝은 청백색으로 이어지며, 밝고 높은 블록일수록 오래된 얼음이다.",
    )
    add_bullet(
        document,
        "블록 높이는 실제 두께가 아니라 연령 범주를 구별하기 위한 정규화된 시각 표현이다.",
    )
    add_bullet(
        document,
        "연도 사이에는 다섯 중간 3D 상태와 8단계 시간 보간을 사용해 60fps로 출력했다.",
    )
    add_bullet(
        document,
        "도입부는 북극 원형 지도에서 1984년으로 확대되며, 마지막은 같은 좌표의 "
        "1984년과 2024년을 나눠 비교한다.",
    )

    add_heading(document, "5. 데이터에서 발견한 변화")
    add_data_table(document)
    source = document.add_paragraph(style="Source")
    source.add_run(
        "주: 화면 지도는 렌더링을 위해 3칸마다 표본화했으나, 위 수치는 "
        "표본화하지 않은 원본 전체 격자로 계산했다."
    )
    add_body(
        document,
        "두 해의 제11주 전체 해빙 범위는 비슷하지만, 그 내부의 시간 구조는 크게 "
        "달랐다. 2024년에는 1984년보다 1년차 얼음의 비중이 커졌고, 오래된 얼음은 "
        "훨씬 적게 남았다. 작품은 면적만으로는 드러나지 않는 이 대비를 시각적 "
        "서사의 중심으로 삼는다.",
    )

    add_heading(document, "6. 전달하고자 하는 환경·생태 메시지")
    add_body(
        document,
        "해빙은 북극 생물이 이용하는 서식 공간이자 바다와 대기 사이에서 열과 "
        "수분이 교환되는 경계다. 오래된 얼음의 감소는 특정 생물종의 변화를 하나의 "
        "원인으로 단정하기 위한 장면이 아니라, 북극 생태계가 의존하는 물리적 환경의 "
        "구조가 과거와 달라지고 있음을 보여주는 신호다.",
    )

    add_heading(document, "7. 과학적 표현 범위와 한계")
    add_bullet(document, "해빙 연령은 실제 두께와 동일한 값이 아니다.")
    add_bullet(
        document,
        "연도 사이의 부드러운 움직임은 시각적 보간이며 관측값이나 예측값이 아니다.",
    )
    add_bullet(
        document,
        "바다에 떠 있는 해빙의 감소를 해수면 상승량으로 환산하지 않았다.",
    )
    add_bullet(
        document,
        "특정 생물종 변화가 오직 해빙 연령 감소 때문에 발생한다고 주장하지 않는다.",
    )

    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    add_heading(document, "8. Q12 제출용 설명문")
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    run = intro.add_run(
        "아래 문단은 온라인 제출 양식의 Q12 항목에 그대로 붙여 넣을 수 있도록 "
        "구성한 설명문이다."
    )
    set_run_font(run, size=10, italic=True, color=MUTED)
    for text in read_q12_paragraphs():
        add_q12_body(document, text)

    add_heading(document, "9. 참고 자료")
    references = (
        "Tschudi, M., Meier, W. N., Stewart, J. S., Fowler, C. & Maslanik, J. "
        "(2019). EASE-Grid Sea Ice Age, Version 4. NASA NSIDC DAAC. "
        "https://doi.org/10.5067/UTAV7490FEPB",
        "National Snow and Ice Data Center. EASE-Grid Sea Ice Age, Version 4 "
        "User Guide. https://nsidc.org/data/nsidc-0611/versions/4",
        "Meier, W. N. et al. (2024). Sea Ice — Arctic Report Card 2024. "
        "https://arctic.noaa.gov/report-card/report-card-2024/sea-ice-2024/",
        "National Snow and Ice Data Center. Why Sea Ice Matters. "
        "https://nsidc.org/learn/parts-cryosphere/sea-ice/why-sea-ice-matters",
    )
    for reference in references:
        paragraph = document.add_paragraph(style="Source")
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(reference)
        set_run_font(run, size=8.5, color=MUTED)

    paragraph = document.add_paragraph(style="Source")
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(
        "라이선스 | 데이터: NSIDC-0611 V4, NASA NSIDC DAAC, "
        "DOI 10.5067/UTAV7490FEPB · 영상 폰트: Wanted Sans 1.0.3, "
        "SIL Open Font License 1.1"
    )
    set_run_font(run, size=8.2, color=MUTED)

    document.core_properties.title = "북극의 기억층 작품 설명서"
    document.core_properties.subject = "2026 극지 빅데이터-인공지능 활용 경진대회"
    document.core_properties.author = "북극의 기억층 출품팀"
    document.core_properties.keywords = "북극, 해빙 연령, NSIDC-0611, 데이터 아트"
    document.save(OUTPUT)
    print(f"wrote {OUTPUT.name.encode('unicode_escape').decode('ascii')}")


if __name__ == "__main__":
    build_document()
